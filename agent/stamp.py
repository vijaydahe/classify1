#!/usr/bin/env python3
"""ClassifyHub document stamping tool.

Stamps a classification label into the header or footer of Office documents and
PDFs, using the organization's policy (placement, font, size, colour) fetched
from the ClassifyHub server. Use it to stamp a file before sharing, or to batch
existing documents.

  python3 stamp.py report.docx --label Confidential
  python3 stamp.py invoice.pdf  --label Restricted --placement footer

Supported: .docx, .pdf  (Word/PDF cover the large majority of documents).
Requires: python-docx (for .docx), reportlab + pypdf (for .pdf):
  pip install python-docx pypdf reportlab
"""
from __future__ import annotations

import argparse
import json
import sys
import urllib.request
from pathlib import Path

CONFIG_PATH = Path(__file__).parent / "config.json"
STATE_PATH = Path(__file__).parent / "state.json"


def fetch_policy() -> dict:
    """Best-effort fetch of the org stamp policy; falls back to sensible defaults."""
    default = {"placement": "footer", "font_name": "Arial", "font_size": 10,
               "color": "#dc2626", "text_template": "CLASSIFICATION: {label}"}
    try:
        cfg = json.loads(CONFIG_PATH.read_text())
        # Stamp policy is exposed per-user; the agent uses its endpoint key only for
        # data APIs, so if a bearer token isn't available we just use defaults.
        token = cfg.get("user_token")
        if not token:
            return default
        req = urllib.request.Request(cfg["server_url"].rstrip("/") + "/api/auth/stamp-policy")
        req.add_header("Authorization", "Bearer " + token)
        with urllib.request.urlopen(req, timeout=15) as r:
            data = json.loads(r.read().decode())
            return {**default, **data}
    except Exception:
        return default


def hex_to_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def stamp_docx(path: Path, text: str, policy: dict) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document(str(path))
    r, g, b = hex_to_rgb(policy["color"])
    for section in doc.sections:
        container = section.header if policy["placement"] == "header" else section.footer
        container.is_linked_to_previous = False
        para = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        # Avoid duplicate stamps on re-run.
        if para.text.strip().startswith(text.split(":")[0]):
            para.clear()
        run = para.add_run(text)
        run.font.name = policy["font_name"]
        run.font.size = Pt(int(policy["font_size"]))
        run.font.bold = True
        run.font.color.rgb = RGBColor(r, g, b)
    doc.save(str(path))


def stamp_pdf(path: Path, text: str, policy: dict) -> None:
    import io
    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    reader = PdfReader(str(path))
    writer = PdfWriter()
    r, g, b = [c / 255 for c in hex_to_rgb(policy["color"])]
    for page in reader.pages:
        w = float(page.mediabox.width); h = float(page.mediabox.height)
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(w, h))
        c.setFillColorRGB(r, g, b)
        font = "Helvetica-Bold"
        c.setFont(font, int(policy["font_size"]))
        y = h - int(policy["font_size"]) - 12 if policy["placement"] == "header" else 12
        c.drawCentredString(w / 2, y, text)
        c.save()
        buf.seek(0)
        overlay = PdfReader(buf).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)
    with open(path, "wb") as f:
        writer.write(f)


def main() -> int:
    ap = argparse.ArgumentParser(description="Stamp a classification label onto a document")
    ap.add_argument("file")
    ap.add_argument("--label", required=True, help="classification label, e.g. Confidential")
    ap.add_argument("--placement", choices=["header", "footer"], help="override policy placement")
    args = ap.parse_args()

    path = Path(args.file)
    if not path.exists():
        print(f"File not found: {path}", file=sys.stderr); return 1

    policy = fetch_policy()
    if args.placement:
        policy["placement"] = args.placement
    text = policy["text_template"].replace("{label}", args.label)

    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            stamp_docx(path, text, policy)
        elif ext == ".pdf":
            stamp_pdf(path, text, policy)
        else:
            print(f"Unsupported file type '{ext}'. Supported: .docx, .pdf", file=sys.stderr)
            return 2
    except ImportError as e:
        print(f"Missing library: {e}. Install with: pip install python-docx pypdf reportlab",
              file=sys.stderr)
        return 3
    print(f"Stamped {path.name} in the {policy['placement']} with: {text}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
